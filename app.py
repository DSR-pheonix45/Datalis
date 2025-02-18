import streamlit as st
import pandas as pd
import numpy as np
import os
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.metrics import accuracy_score, mean_squared_error, classification_report, confusion_matrix
from sklearn.preprocessing import LabelEncoder
from langchain_groq import ChatGroq
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
import joblib
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from scipy.stats import zscore
import gspread
import pandasql as ps

def display_intro():
    st.title("Welcome to Datalis")
    st.markdown("""
    **Datalis** is a comprehensive data analysis tool designed to help you clean, transform, visualize, and gain insights from your data efficiently. Here's a quick guide on how to navigate through the app:

    - **Upload Data**: Start by uploading your CSV or Excel files. You can drag and drop multiple files, which will be organized into folders for easy access.
    - **Data Cleaning**: Use this section to handle missing values, remove duplicates, and perform other cleaning operations on your dataset.
    - **Data Transformation**: Transform your data by renaming columns, converting data types, and more.
    - **Data Visualization**: Create various types of charts to visualize your data, including histograms, line charts, and scatter plots.
    - **AI Chat Platform**: Interact with our AI to get insights and answers about your data.
    - **Export Report**: Generate and download reports in Word format based on your analysis.

    Navigate through these sections using the sidebar on the left. Let's get started by uploading your data!
    """)

def upload_files():
    display_intro()
    st.header("Upload Data")
    uploaded_file = st.file_uploader("Drag and drop a CSV or Excel file here", type=["csv", "xlsx", "xls"], accept_multiple_files=False)

    if uploaded_file:
        upload_dir = "uploaded_files"
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)

        file_path = os.path.join(upload_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success(f"Uploaded {uploaded_file.name}")

        st.write("### Uploaded File")
        st.write(uploaded_file.name)

        if uploaded_file.name.endswith(".csv"):
            st.session_state.df = pd.read_csv(file_path)
        elif uploaded_file.name.endswith((".xlsx", ".xls")):
            st.session_state.df = pd.read_excel(file_path)

        st.session_state.uploaded_file = uploaded_file.name
        st.write("### Data Preview")
        st.dataframe(st.session_state.df, height=300, use_container_width=True)

    # Add Next button
    if st.button("Next", key="next_to_cleaning"):
        st.session_state.page = "Data Cleaning"

def select_file():
    if 'uploaded_file' in st.session_state:
        file_path = os.path.join("uploaded_files", st.session_state.uploaded_file)
        if st.session_state.uploaded_file.endswith(".csv"):
            st.session_state.df = pd.read_csv(file_path)
        elif st.session_state.uploaded_file.endswith((".xlsx", ".xls")):
            st.session_state.df = pd.read_excel(file_path)

# Function to select operations for cleaning or transformation
def select_operations(operation_type):
    st.warning(f"Select {operation_type.capitalize()} Operations")
    operations = {}

    if operation_type == "cleaning":
        operations["handle_missing_values"] = st.checkbox("Handle Missing Values", value=True)
        operations["remove_duplicates"] = st.checkbox("Remove Duplicates", value=True)
        operations["handle_outliers"] = st.checkbox("Handle Outliers (Z-score)", value=True)
        operations["handle_negative_values"] = st.checkbox("Handle Negative Values", value=True)

    elif operation_type == "transformation":
        operations["rename_columns"] = st.checkbox("Rename Columns (to lowercase with underscores)", value=True)
        operations["convert_data_types"] = st.checkbox("Convert Object Columns to Category", value=True)
        operations["feature_engineering"] = st.checkbox("Feature Engineering (Extract Date Parts)", value=True)
        operations["bin_numeric_columns"] = st.checkbox("Bin Numeric Columns", value=True)
        operations["normalize_numeric_columns"] = st.checkbox("Normalize Numeric Columns", value=True)
        operations["one_hot_encode"] = st.checkbox("One-Hot Encode Categorical Columns", value=True)

    return operations

# Advanced Data Cleaning
def clean_data(df, operations):
    if operations["handle_missing_values"]:
        df = df.ffill().bfill()

    if operations["remove_duplicates"]:
        df = df.drop_duplicates()

    if operations["handle_outliers"]:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            z_scores = zscore(df[col])
            abs_z_scores = np.abs(z_scores)
            filtered_entries = (abs_z_scores < 3)
            df = df[filtered_entries]

    if operations["handle_negative_values"]:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            if (df[col] < 0).any(): 
                df[col] = df[col].apply(lambda x: np.nan if x < 0 else x)
                df[col] = df[col].ffill().bfill()

    return df

# Advanced Data Transformation
def transform_data(df, operations):
    if operations["rename_columns"]:
        df.columns = [col.replace(" ", "_").lower() for col in df.columns]

    if operations["convert_data_types"]:
        for col in df.columns:
            if df[col].dtype == 'object':
                df[col] = df[col].astype('category')

    if operations["feature_engineering"] and 'date' in df.columns:
        df['year'] = pd.to_datetime(df['date']).dt.year
        df['month'] = pd.to_datetime(df['date']).dt.month
        df['day'] = pd.to_datetime(df['date']).dt.day

    if operations["bin_numeric_columns"]:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            if df[col].nunique() > 10: 
                df[f'{col}_binned'] = pd.cut(df[col], bins=5, labels=False)

    if operations["normalize_numeric_columns"]:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            if df[col].nunique() > 1:
                df[col] = (df[col] - df[col].mean()) / df[col].std()

    if operations["one_hot_encode"]:
        categorical_cols = df.select_dtypes(include=['category']).columns
        df = pd.get_dummies(df, columns=categorical_cols, drop_first=True)

    return df

# Generate Word report
def generate_word_report(summary, graphs, filename):
    doc = Document()
    doc.add_heading("Data Summary Report", level=1)
    doc.add_paragraph("Data Summary:")
    doc.add_paragraph(summary)

    for i, fig in enumerate(graphs):
        img_filename = f"temp_chart_{i}.png"
        fig.write_image(img_filename)
        doc.add_paragraph(fig.layout.title.text)
        doc.add_picture(img_filename, width=docx.shared.Inches(6))
        os.remove(img_filename)

    doc.save(filename)

# Groq AI bot - Modified for human-like responses
def get_groq_response(user_prompt, df, data_scope="full"):
    try:
        llm = ChatGroq(
            temperature=0.7,
            groq_api_key=os.getenv("GROQ_API_KEY"),
            model_name="llama-3.2-3b-preview"
        )

        total_rows = len(df)
        if total_rows > 500:
            disclaimer = (
                "⚠️ Your dataset is large (more than 500 rows). "
                "Would you like insights from:\n"
                "1️⃣ Entire dataset\n"
                "2️⃣ First half\n"
                "3️⃣ Last half\n\n"
                "By default, analyzing the full dataset."
            )
            print(disclaimer)  # Log disclaimer for user awareness

        # Select data scope based on user choice
        if data_scope == "first_half":
            df = df.iloc[: total_rows // 2]
        elif data_scope == "last_half":
            df = df.iloc[total_rows // 2 :]

        prompt = (
            f"The dataset has these columns: {df.columns.tolist()}.\n\n"
            f"Here's a sneak peek ({len(df)} rows considered):\n{df.to_string(index=False)}\n\n"
            f"User Query: {user_prompt}\n\n"
            "Provide a **one-statement summary** first, then list key **highlighted data points** before diving into detailed insights.\n"
            "Ensure the response is 90% based on the dataset and 10% creatively related to its domain.\n"
            "Give crisp, data-driven insights like a professional data analyst.\n"
            "Avoid any code in the response."
        )

        response = llm.invoke(prompt)
        return response.content

    except Exception as e:
        if "tokens" in str(e).lower():  # Token exhaustion check
            return "🚨 Token limit reached! Subscribe to premium for unlimited insights."
        return f"❌ Error: {str(e)}"

# Page for Data Cleaning
def cleaning_page():
    st.header("Data Cleaning")
    select_file()
    if st.session_state.df is not None:
        st.write("### Data Preview")

        # Toggle button for editing
        edit_mode = st.checkbox("Enable Editing", value=False)

        if edit_mode:
            # Display editable DataFrame
            edited_df = st.data_editor(st.session_state.df, height=300, use_container_width=True)
            # Update the session state with the edited DataFrame
            st.session_state.df = edited_df
        else:
            # Display non-editable DataFrame
            st.dataframe(st.session_state.df, height=300, use_container_width=True)

        cleaning_operations = select_operations("cleaning")

        if st.button("Clean Data"):
            st.session_state.df = clean_data(st.session_state.df.copy(), cleaning_operations) 
            st.write("Data cleaned successfully!")
            st.write("### Cleaned Data Preview")
            st.dataframe(st.session_state.df, height=300, use_container_width=True)

            csv_filename = "cleaned_data.csv"
            st.download_button(
                label="Download Cleaned CSV",
                data=st.session_state.df.to_csv(index=False).encode('utf-8'),
                file_name=csv_filename,
                mime="text/csv"
            )

        # Add Next button
        if st.button("Next", key="next_to_transformation"):
            st.session_state.page = "Data Transformation"
    else:
        st.warning("No data available. Please upload a CSV or Excel file.")

# Page for Data Transformation
def transformation_page():
    st.header("Data Transformation")
    select_file()
    if st.session_state.df is not None:
        st.write("### Data Preview")

        # Toggle button for editing
        edit_mode = st.checkbox("Enable Editing", value=False)

        if edit_mode:
            # Display editable DataFrame
            edited_df = st.data_editor(st.session_state.df, height=300, use_container_width=True)
            # Update the session state with the edited DataFrame
            st.session_state.df = edited_df
        else:
            # Display non-editable DataFrame
            st.dataframe(st.session_state.df, height=300, use_container_width=True)

        transformation_operations = select_operations("transformation")

        if st.button("Transform Data"):
            st.session_state.df = transform_data(st.session_state.df.copy(), transformation_operations) 
            st.write("Data transformed successfully!")
            st.write("### Transformed Data Preview")
            st.dataframe(st.session_state.df, height=300, use_container_width=True)

            csv_filename = "transformed_data.csv"
            st.download_button(
                label="Download Transformed CSV",
                data=st.session_state.df.to_csv(index=False).encode('utf-8'),
                file_name=csv_filename,
                mime="text/csv"
            )

        # Add Next button
        if st.button("Next", key="next_to_visualization"):
            st.session_state.page = "Data Visualization"
    else:
        st.warning("No data available. Please upload a CSV or Excel file.")

# Page for Visualization
def visualization_page():
    st.header("Data Visualization")
    select_file()
    if st.session_state.df is not None:
        df = st.session_state.df

        chart_type = st.selectbox("Select chart type", ["Histogram", "Line Chart", "Bar Chart", 
                                                        "Scatter Plot", "Box Plot", "Donut Chart", "Heatmap"])
        
        # Store chart parameters in session state
        if 'chart_params' not in st.session_state:
            st.session_state.chart_params = {}
        
        if chart_type == "Histogram":
            st.session_state.chart_params['column'] = st.selectbox("Select a column to visualize", df.columns)
        elif chart_type == "Line Chart":
            st.session_state.chart_params['x_column'] = st.selectbox("Select X-axis column", df.columns)
            st.session_state.chart_params['y_column'] = st.selectbox("Select Y-axis column", df.columns)
        elif chart_type == "Bar Chart": 
            st.session_state.chart_params['x_column'] = st.selectbox("Select X-axis column", df.columns)
            st.session_state.chart_params['y_column'] = st.selectbox("Select Y-axis column", df.columns)
        elif chart_type == "Scatter Plot":
            st.session_state.chart_params['x_column'] = st.selectbox("Select X-axis column", df.columns)
            st.session_state.chart_params['y_column'] = st.selectbox("Select Y-axis column", df.columns)
            st.session_state.chart_params['color_column'] = st.selectbox("Select column for color (optional)", [None] + df.columns.tolist())
        elif chart_type == "Box Plot":
            st.session_state.chart_params['column'] = st.selectbox("Select a column to visualize", df.columns)
            st.session_state.chart_params['group_column'] = st.selectbox("Select column for grouping (optional)", [None] + df.columns.tolist())
        elif chart_type == "Donut Chart":
            st.session_state.chart_params['column'] = st.selectbox("Select a column for the donut chart", df.columns)
        elif chart_type == "Heatmap":
            st.session_state.chart_params['x_column'] = st.selectbox("Select X-axis column for heatmap", df.columns)
            st.session_state.chart_params['y_column'] = st.selectbox("Select Y-axis column for heatmap", df.columns)
            st.session_state.chart_params['z_column'] = st.selectbox("Select value column for heatmap", df.columns)

        selected_graphs = []

        # Button to generate the chart
        if st.button("Generate Chart"):
            if chart_type == "Histogram":
                column = st.session_state.chart_params.get('column')
                if pd.api.types.is_numeric_dtype(df[column]):
                    fig = px.histogram(df, x=column, nbins=50, title=f"Distribution of {column}",
                                       labels={'x': 'Value', 'y': 'Frequency'}) 
                    st.plotly_chart(fig)
                    selected_graphs.append(fig)
            elif chart_type == "Line Chart":
                x_column = st.session_state.chart_params.get('x_column')
                y_column = st.session_state.chart_params.get('y_column')
                fig = px.line(df, x=x_column, y=y_column, title=f"Line Chart of {x_column} vs {y_column}", 
                              labels={'x': x_column, 'y': y_column})
                st.plotly_chart(fig)
                selected_graphs.append(fig)
            elif chart_type == "Bar Chart": 
                x_column = st.session_state.chart_params.get('x_column')
                y_column = st.session_state.chart_params.get('y_column')
                fig = px.bar(df, x=x_column, y=y_column, title=f"Bar Chart of {x_column} vs {y_column}",
                             labels={'x': x_column, 'y': y_column})
                st.plotly_chart(fig)
                selected_graphs.append(fig)
            elif chart_type == "Scatter Plot":
                x_column = st.session_state.chart_params.get('x_column')
                y_column = st.session_state.chart_params.get('y_column')
                color_column = st.session_state.chart_params.get('color_column')
                fig = px.scatter(df, x=x_column, y=y_column, color=color_column, 
                                 title=f"Scatter Plot of {x_column} vs {y_column}",
                                 labels={'x': x_column, 'y': y_column, 'color': color_column})
                st.plotly_chart(fig)
                selected_graphs.append(fig)
            elif chart_type == "Box Plot":
                column = st.session_state.chart_params.get('column')
                group_column = st.session_state.chart_params.get('group_column')
                if pd.api.types.is_numeric_dtype(df[column]):
                    fig = px.box(df, y=column, color=group_column, title=f"Box Plot of {column}",
                                 labels={'y': column, 'color': group_column})
                    st.plotly_chart(fig)
                    selected_graphs.append(fig)
            elif chart_type == "Donut Chart":
                column = st.session_state.chart_params.get('column')
                fig = px.pie(df, names=column, title=f"Donut Chart of {column}", hole=0.4,
                             labels={'names': column})
                st.plotly_chart(fig)
                selected_graphs.append(fig)
            elif chart_type == "Heatmap":
                x_column = st.session_state.chart_params.get('x_column')
                y_column = st.session_state.chart_params.get('y_column')
                z_column = st.session_state.chart_params.get('z_column')
                fig = px.imshow(df.pivot(index=y_column, columns=x_column, values=z_column), 
                                title=f"Heatmap of {z_column} by {x_column} and {y_column}",
                                labels=dict(x=x_column, y=y_column, color=z_column))
                st.plotly_chart(fig)
                selected_graphs.append(fig)

        if selected_graphs:
            st.write("### Select Graphs to Include in Report")
            for i, graph in enumerate(selected_graphs):
                if st.checkbox(f"Include {graph.layout.title.text} in report", value=True, key=f"checkbox_{i}"):
                    if 'selected_graphs' not in st.session_state:
                        st.session_state.selected_graphs = []
                    st.session_state.selected_graphs.append(graph)

        # Generate and store the summary in session state
        if st.button("Generate Summary"):
            with st.spinner('Generating data summary...'):
                st.session_state.summary = get_groq_response("Provide a summary of the dataset.", st.session_state.df)
                st.success("Data summary generated!")

        # Add Next button
        if st.button("Next", key="next_to_ai_chat"):
            st.session_state.page = "AI Chat Platform"
    else:
        st.warning("No data available. Please upload a CSV or Excel file.")

# Page for AI Chat Platform
def ai_chat_page():
    st.header("AI Chat Platform")
    
    # Ensure uploaded_file is initialized
    if 'uploaded_file' not in st.session_state:
        st.session_state.uploaded_file = None

    # Display uploaded file information
    if st.session_state.uploaded_file is not None:
        st.success(f"Uploaded File: {st.session_state.uploaded_file}")
        st.write("### Data Preview")
        st.dataframe(st.session_state.df.head(20), height=200, use_container_width=True)

        # User query input
        user_query = st.text_input("Ask the Dabby anything about the data:", key="user_query", label_visibility="collapsed")
        
        # Initialize chat history in session state
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []

        # Process user query
        if st.button("Send") or (user_query and st.session_state.get("enter_pressed", False)):
            response = get_groq_response(user_query, st.session_state.df)
            st.write("### Datalis Dabby Suggests")
            st.write(response)
            
            # Record the query and response in chat history
            st.session_state.chat_history.append({"query": user_query, "response": response})
            st.session_state["enter_pressed"] = False

        if st.session_state.get("user_query") and st.session_state.get("user_query") != "":
            st.session_state["enter_pressed"] = True

        # Add Next button
        if st.button("Next", key="next_to_export"):
            st.session_state.page = "Export Report"
    else:
        st.warning("No data available. Please upload a CSV or Excel file.")


# Page for Export Report
@st.cache_resource
def render_chart(_fig, _key):
    st.plotly_chart(_fig, key=_key, use_container_width=True) 

def export_report_page():
    st.header("Export Report")
    select_file()
    if st.session_state.df is not None:
        # Access pre-generated summary
        summary = st.session_state.summary 

        st.write("### Data Summary")
        st.write(summary)

        st.write("### Selected Graphs for Report")
        if 'selected_graphs' in st.session_state and st.session_state.selected_graphs:
            for i, graph in enumerate(st.session_state.selected_graphs):
                render_chart(go.Figure(graph), f"chart_{i}")

        st.write("### Export Options")
        
        if st.button("Generate Word Report"):
            with st.spinner('Generating Word report...'):
                report_filename = "data_summary_report.docx"
                generate_word_report(summary, st.session_state.selected_graphs, report_filename)
                with open(report_filename, "rb") as f:
                    st.download_button(
                        label="Download Word Report",
                        data=f,
                        file_name=report_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.success("Word report generated successfully!")
    else:
        st.warning("No data available. Please upload a CSV or Excel file.")

# Main App Functionality
def main():
    st.markdown(
        """
        <style>
        /* General Styles */
        body {
            font-family: 'Arial', sans-serif;
            background-color: #f8f8f8; 
            color: #333;
            margin: 0;
            padding: 0;
            line-height: 1.6; 
        }

        .container { 
            max-width: 1200px;
            margin: 20px auto;
            padding: 20px;
            background-color: #fff; 
            box-shadow: 0px 2px 5px rgba(0, 0, 0, 0.1); 
            border-radius: 8px;
        }

        header { 
            background: linear-gradient(135deg, #667eea, #764ba2); 
            padding: 30px 0; 
            text-align: center;
            color: white; 
        }

        header h1 { 
            font-size: 2rem; 
            margin-bottom: 0.5rem;
        }

        /* Button Styling */
        .stButton>button { 
            background-color: #4CAF50; /* Green */ 
            color: white;
            border: none;
            padding: 10px 20px; 
            font-size: 1rem;
            border-radius: 5px;
            transition: background-color 0.3s ease; 
            box-shadow: 0px 2px 4px rgba(0, 0, 0, 0.1); 
        }

        .stButton>button:hover {
            background-color: #45a049;  /* Slightly Darker Green */
            cursor: pointer; 
        }

        /* Input Fields & Text Areas */
        .stTextInput input, 
        .stTextInput textarea { 
            border: 1px solid #ddd; 
            border-radius: 5px;
            padding: 10px;
            width: calc(100% - 22px);  
            font-size: 1rem; 
            box-sizing: border-box;
        }

        .stTextInput input:focus, 
        .stTextInput textarea:focus {
            outline: none; 
            border-color: #4CAF50;  
            box-shadow: 0px 0px 5px rgba(76, 175, 80, 0.3); 
        }

        /* Sidebar Styles - Enhanced Matte Look */
        .sidebar .sidebar-content {
            background-color: #4d4d4d; /* Darker Gray */
            padding: 20px;
            border-radius: 8px;
        }

        .sidebar .stButton>button { 
            background-color: #e7e7e7; 
            color: #333;
            border: none;
            padding: 10px 15px;
            width: 100%;
            margin-bottom: 10px; 
            text-align: left; 
            border-radius: 5px;
            transition: all 0.3s ease; 
        }

        .sidebar .stButton>button:hover { 
            background-color: #d4d4d4; 
        }

        .sidebar .stButton>button:before { 
            font-family: "Font Awesome 5 Free"; 
            margin-right: 10px;  
            display: inline-block;
            vertical-align: middle; 
        }

        /* Font Awesome Icons for Sidebar */
        <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.4/css/all.min.css"> 

        .sidebar .stButton[data-testid="upload_data"]>button:before {
           content: "\f093"; /* Upload Icon */
        }
        .sidebar .stButton[data-testid="data_cleaning"]>button:before {
           content: "\f1ea"; /* Broom (Cleaning) Icon */
        }
        .sidebar .stButton[data-testid="data_transformation"]>button:before {
           content: "\f0ec"; /* Exchange (Transformation) Icon */
        }
        .sidebar .stButton[data-testid="data_visualization"]>button:before {
           content: "\f57d"; /* Chart Bar Icon */
        } 
        .sidebar .stButton[data-testid="ai_chat_platform"]>button:before {
           content: "\f086"; /* Comment Dots Icon */
        } 
        .sidebar .stButton[data-testid="export_report"]>button:before {
           content: "\f019"; /* Download Icon */
        }

        /* Footer - Optional, if used */
        footer {
            background: #2c3e50; 
            padding: 20px 0;
            text-align: center;
            color: white;
            margin-top: 40px; 
        }
        </style> 
        """,
        unsafe_allow_html=True,
    )

    if 'df' not in st.session_state:
        st.session_state.df = None

    if 'selected_graphs' not in st.session_state:
        st.session_state.selected_graphs = []

    if 'summary' not in st.session_state:
        st.session_state.summary = None

    # Directly go to app functionality 
    st.title("Datalis")

    st.sidebar.title("Navigation")
    if st.sidebar.button("Upload Data", key="upload_data"):
        st.session_state.page = "Upload Data"
    if st.sidebar.button("Data Cleaning", key="data_cleaning"):
        st.session_state.page = "Data Cleaning"
    if st.sidebar.button("Data Transformation", key="data_transformation"):
        st.session_state.page = "Data Transformation"
    if st.sidebar.button("Data Visualization", key="data_visualization"):
        st.session_state.page = "Data Visualization"
    if st.sidebar.button("AI Chat Platform", key="ai_chat_platform"):
        st.session_state.page = "AI Chat Platform"
    if st.sidebar.button("Export Report", key="export_report"):
        st.session_state.page = "Export Report"

    if 'page' not in st.session_state:
        st.session_state.page = "Upload Data"

    if st.session_state.page == "Upload Data":
        upload_files()

    elif st.session_state.page == "Data Cleaning":
        cleaning_page()

    elif st.session_state.page == "Data Transformation":
        transformation_page()

    elif st.session_state.page == "Data Visualization":
        visualization_page()

    elif st.session_state.page == "AI Chat Platform":
        ai_chat_page()

    elif st.session_state.page == "Export Report":
        export_report_page()

if __name__ == '__main__':
    main()

